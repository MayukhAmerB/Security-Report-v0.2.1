import mysql.connector
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def fetch_vulnerability_data(vulnerability_id):
    # Connect to the MySQL database
    connection = mysql.connector.connect(
        host='your_host',
        user='your_user',
        password='your_password',
        database='your_database'
    )
    cursor = connection.cursor(dictionary=True)

    # Fetch vulnerability data
    query = "SELECT * FROM vulnerabilities WHERE id = %s"
    cursor.execute(query, (vulnerability_id,))
    result = cursor.fetchone()

    cursor.close()
    connection.close()

    return result


def create_document(vulnerability_data):
    # Create a new Document
    doc = Document()

    # Add Title
    title = doc.add_heading('Vulnerability Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Define fields and their corresponding keys in the database
    fields = [
        ("New or Repeat observation", "new_or_repeat"),
        ("Affected Asset", "affected_asset"),
        ("Vulnerability Title", "vulnerability_title"),
        ("Severity", "severity"),
        ("Impact", "impact"),
        ("Likelihood", "likelihood"),
        ("Business Impact", "business_impact"),
        ("Detailed Observation", "detailed_observation"),
        ("CVE/CWE", "cve_cwe"),
        ("Proof of Concept", "proof_of_concept"),
        ("Vulnerable Points", "vulnerable_points"),
        ("Remediation", "remediation")  # New Field
    ]

    for field_name, db_key in fields:
        doc.add_heading(field_name, level=2)
        if vulnerability_data and db_key in vulnerability_data:
            doc.add_paragraph(vulnerability_data[db_key])
        else:
            doc.add_paragraph("N/A")

    # Save the document
    doc.save('Vulnerability_Report.docx')


if __name__ == "__main__":
    # Example vulnerability ID to fetch from the database
    vulnerability_id = 1
    vulnerability_data = fetch_vulnerability_data(vulnerability_id)
    create_document(vulnerability_data)
