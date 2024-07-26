from flask import Flask, request, jsonify, render_template, redirect, url_for, send_file
import io
import os
import mysql.connector
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from flask_session import Session
from threading import Timer
from datetime import timedelta
app = Flask(__name__)

# Configure upload folder
app.config['UPLOAD_FOLDER'] = './static/vuln_images'
app.config['REPORT_UPLOAD_FOLDER'] = './uploads'
app.config['PDF_UPLOAD_FOLDER'] = './uploads/PDF_UPLOAD_FOLDER'
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])
if not os.path.exists(app.config['REPORT_UPLOAD_FOLDER']):
    os.makedirs(app.config['REPORT_UPLOAD_FOLDER'])
app.config['PDF_UPLOAD_FOLDER'] = './uploads/PDF_UPLOAD_FOLDER'
if not os.path.exists(app.config['PDF_UPLOAD_FOLDER']):
    os.makedirs(app.config['PDF_UPLOAD_FOLDER'])



# Session configuration
app.config['SECRET_KEY'] = 'your_secret_key'
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_USE_SIGNER'] = True
app.config['SESSION_KEY_PREFIX'] = 'sess:'

Session(app)

# In-memory storage for vulnerabilities and document data
default_data = {
    'selected_vulnerabilities': [],
    'document_data': {},
    'document_preparation_data': {},
    'document_control_data': {},
    'document_change_history_data': [],
    'document_distribution_list_data': [],
    'contents_data': [],
    'engagement_scope_data': [],
    'auditing_team_data': [],
    'audit_timeline_data': [],
    'audit_methodology_criteria_data': {},
    'owasp_data': {},
    'sans_data': [],
    'tools_data' : [],
    'title_data' : {},
    'disclaimer_data' :{},
}

# Initialize session data
def initialize_session():
    for key, value in default_data.items():
        Session[key] = value

@app.before_request
def make_session_permanent():
    Session.permanent = True
    app.permanent_session_lifetime = timedelta(minutes=30)  # 30 minutes session lifetime

# Background task to clear session data
def clear_session():
    with app.app_context():
        initialize_session()
        print("Session data cleared")

def set_session_timeout():
    Timer(1800, clear_session).start()  # 1800 seconds = 30 minutes
# Database connection
def get_db_connection():
    connection = mysql.connector.connect(
        host="localhost",
        user="root",
        password="root",
        auth_plugin='mysql_native_password',
        database="security_assessment_tool"
    )
    return connection

# In-memory storage for vulnerabilities and document data
selected_vulnerabilities = []
document_data = {}
document_preparation_data = {}
document_control_data = {}
document_change_history_data = []
document_distribution_list_data = []
contents_data = []
engagement_scope_data= []
auditing_team_data=[]
audit_timeline_data = []
audit_methodology_criteria_data={}
owasp_data = {}
sans_data = []
tools_data = []
title_data = {}
disclaimer_data = {
    'acronym': '',
    'disclaimer': ''
}
appendix_data = {
    'risk_ranking_approach': '',
    'likelihood': '',
    'impact': ''
}


@app.route('/document_preparation', methods=['GET', 'POST'])
def document_preparation():
    if request.method == 'POST':
        # Handle data submission
        data = request.json
        # Save the data (In real application, you might save this to a database)
        global document_preparation_data
        document_preparation_data = data
        return jsonify({'success': True, 'message': 'Data saved successfully'})

    # Handle data retrieval for GET request
    return render_template('Document Preparation.html', document_data=document_preparation_data)

@app.route('/save_document_preparation', methods=['POST'])
def save_document_preparation():
    global document_preparation_data
    document_preparation_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully'})

@app.route('/get_document_preparation_data', methods=['GET'])
def get_document_preparation_data():
    return jsonify(document_preparation_data)

@app.route('/save_document_control', methods=['POST'])
def save_document_control():
    global document_control_data
    document_control_data = request.json
    # Perform any backend processing or printing here
    return jsonify({'success': True, 'message': 'Data saved successfully'})

@app.route('/document_change_history', methods=['GET', 'POST'])
def document_change_history():
    if request.method == 'POST':
        # Handle data submission
        global document_change_history_data
        document_change_history_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})

    # Handle data retrieval for GET request
    return render_template('Document Change History.html', document_data=document_change_history_data)


@app.route('/save_document_change_history', methods=['POST'])
def save_document_change_history():
    global document_change_history_data
    document_change_history_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})


@app.route('/get_document_change_history_data', methods=['GET'])
def get_document_change_history_data():
    return jsonify(document_change_history_data)



@app.route('/')
def index():
    return render_template('index.html')

@app.route('/show_vulnerabilities', methods=['GET'])
def show_vulnerabilities():
    category = request.args.get('category', 'N/A')
    return render_template('vulnerabilities.html', vulnerabilities=selected_vulnerabilities, category=category)


@app.route('/get_vulnerability', methods=['POST'])
def get_vulnerability():
    global selected_vulnerabilities
    data = request.get_json()
    vulnerabilities = data.get('vulnerabilities', [])
    category = data.get('category')

    if not vulnerabilities or not category:
        return jsonify({'message': 'Invalid data', 'success': False}), 400

    if category == 'WASA':
        connection = get_db_connection()
        cursor = connection.cursor()
        selected_vulnerabilities = []
        try:
            for vuln in vulnerabilities:
                cursor.execute(
                    """
                    SELECT Business_Impact, detailed_observation, Recommendation, Reference, 
                           CONCAT(CVE_No, ' / ', CWE_No) AS CVE_CWE_No 
                    FROM vulnerabilities 
                    WHERE Vulnerability_title = %s
                    """, (vuln,)
                )

                result = cursor.fetchone()

                if result:
                    selected_vulnerabilities.append({
                        'id': len(selected_vulnerabilities),
                        'name': vuln,
                        'business_impact': result[0],
                        'detailed_observation': result[1],
                        'recommendation': result[2],
                        'reference': result[3],
                        'cve_cwe': result[4],
                        'pocs': []
                    })
                else:
                    return jsonify({'message': f'Vulnerability {vuln} not found', 'success': False}), 404

            if selected_vulnerabilities:
                return jsonify({'success': True, 'redirect_url': url_for('show_vulnerabilities', category=category)})
            else:
                return jsonify({'message': 'No vulnerabilities selected', 'success': False}), 400

        except mysql.connector.Error as err:
            return jsonify({'message': str(err), 'success': False}), 500

        finally:
            cursor.close()
            connection.close()
    else:
        return jsonify({'message': 'Invalid category', 'success': False}), 400

@app.route('/edit_vulnerability/<int:vuln_id>', methods=['GET', 'POST'])
def edit_vulnerability(vuln_id):
    global selected_vulnerabilities
    if request.method == 'POST':
        updated_data = request.form.to_dict()
        updated_data['id'] = vuln_id

        # Handle file upload for Proof of Concept
        poc_files = []
        poc_descriptions = []
        for key in request.files:
            if key.startswith('poc_file_'):
                file = request.files[key]
                if file.filename != '':
                    filename = secure_filename(file.filename)
                    file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    poc_files.append(filename)
        for key in request.form:
            if key.startswith('poc_description_'):
                poc_descriptions.append(request.form[key])

        new_pocs = list(zip(poc_files, poc_descriptions))

        for vuln in selected_vulnerabilities:
            if vuln['id'] == vuln_id:
                if 'pocs' not in vuln:
                    vuln['pocs'] = []
                vuln['pocs'].extend(new_pocs)  # Merge old and new PoCs
                vuln.update(updated_data)
        return redirect(url_for('show_vulnerabilities', category=request.args.get('category', 'N/A')))

    vulnerability = next((v for v in selected_vulnerabilities if v['id'] == vuln_id), None)
    return render_template('edit_vulnerability.html', vulnerability=vulnerability)

@app.route('/delete_poc/<filename>', methods=['DELETE'])
def delete_poc(filename):
    global selected_vulnerabilities
    try:
        # Delete the file from the upload directory
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.exists(file_path):
            os.remove(file_path)
        else:
            return jsonify({'success': False, 'message': 'File not found'}), 404

        # Remove the POC entry from the selected vulnerabilities list
        for vuln in selected_vulnerabilities:
            vuln['pocs'] = [poc for poc in vuln['pocs'] if poc[0] != filename]

        return jsonify({'success': True, 'message': 'Proof of Concept deleted successfully'})
    except Exception as e:
        app.logger.error(f'Error deleting Proof of Concept: {e}')
        return jsonify({'success': False, 'message': 'Error deleting Proof of Concept'}), 500

@app.route('/document_control')
def document_control():
    return render_template('document_control.html')

@app.route('/download_report')
def download_report():
    global selected_vulnerabilities, document_data
    format = request.args.get('format', 'word')

    if format == 'word':
        response = download_report_word(selected_vulnerabilities)
    elif format == 'text':
        response = download_report_text(selected_vulnerabilities)
    else:
        return jsonify({'error': 'Invalid format selected.'}), 400

    return response

def download_report_text(selected_vulnerabilities):
    report_content = "Selected Vulnerabilities Report\n\n"
    for vuln in selected_vulnerabilities:
        report_content += f"Affected Asset: {vuln.get('assets', 'N/A')}\n"
        report_content += f"Vulnerability title: {vuln.get('name', 'N/A')}\n"
        report_content += f"Severity: {vuln.get('severity', 'N/A')}\n"
        report_content += "Business Impact:\n"
        report_content += f"{vuln.get('business_impact', 'N/A')}\n"
        report_content += "Detailed observation:\n"
        report_content += f"{vuln.get('detailed_observation', 'N/A')}\n"
        report_content += f"CVE/CWE: {vuln.get('cve_cwe', 'N/A')}\n"
        report_content += "Proof of Concept:\n"
        for poc in vuln.get('pocs', []):
            report_content += f"  - {poc[1]} (Image: {poc[0]})\n"
        report_content += "Vulnerable Points:\n"
        report_content += f"{vuln.get('vulnerable_points', 'N/A')}\n"
        report_content += "Recommendation:\n"
        report_content += f"{vuln.get('recommendation', 'N/A')}\n"
        report_content += f"Reference: {vuln.get('reference', 'N/A')}\n"
        report_content += f"New or Repeat observation: {vuln.get('new_repeat_observation', 'N/A')}\n\n"

    report_file = io.BytesIO()
    report_file.write(report_content.encode('utf-8'))
    report_file.seek(0)

    return send_file(report_file, mimetype='text/plain', as_attachment=True, download_name='vulnerabilities_report.txt')

def set_table_border(table):
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')

    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_element = OxmlElement(f'w:{border}')
        border_element.set(qn('w:val'), 'single')
        border_element.set(qn('w:sz'), '8')
        border_element.set(qn('w:space'), '0')
        border_element.set(qn('w:color'), '000000')
        tblBorders.append(border_element)

    tbl.tblPr.append(tblBorders)
def set_table_header_style_owasp(cell, text):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Set background color
    cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
    cell._element.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), 'a0c4e4')
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Set font color to white

def set_table_header_style_sans(cell, text):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Set background color
    cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
    cell._element.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), 'a0c4e4')
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Set font color to white
def set_table_header_style_AA(cell, text):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Set background color
    cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
    cell._element.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), 'a0c4e4')
    run.font.color.rgb =RGBColor(0x00, 0x00, 0x00)  # Set font color to white


def set_vulnerability_table_cell_text(cell, heading, text, font_size, bold_heading=False):
    # Add the header text
    heading_paragraph = cell.add_paragraph(heading)
    if bold_heading:
        for run in heading_paragraph.runs:
            run.font.bold = True

    # Apply formatting to the header
    for run in heading_paragraph.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(font_size)

    # Add an empty paragraph to create a line gap between the header and the text
    if text:
        cell.add_paragraph()  # Adding an empty paragraph for line gap
        text_paragraph = cell.add_paragraph(text)
        for run in text_paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(font_size)

def set_document_cell_text_engagement(cell, text):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)  # Adjust as needed
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def set_header_style_engagement(cell, text, font_size=12):
    """
    Set the header text in a table cell with specific styling and vertical orientation.

    :param cell: The table cell where the header text will be set.
    :param text: The header text.
    :param font_size: The font size for the header text.
    """
    cell.text = ""
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(font_size)
    font.bold = True
    font.color.rgb =RGBColor(0x00, 0x00, 0x00)  # Set font color to #FFFFFF
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'a0c4e4')  # Set background color to #31473A
    cell._element.get_or_add_tcPr().append(shading_elm)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set vertical text direction
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    text_direction = OxmlElement('w:textDirection')
    text_direction.set(qn('w:val'), 'btLr')
    tcPr.append(text_direction)


def set_header_style_Content(cell, text, font_size=12):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(font_size)
    font.bold = True
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Set font color to #FFFFFF
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'a0c4e4')  # Set background color to #31473A
    cell._element.get_or_add_tcPr().append(shading_elm)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def set_header_style_Audit(cell, text, font_size=12):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(font_size)
    font.bold = True
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Set font color to #FFFFFF
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'a0c4e4')
    cell._element.get_or_add_tcPr().append(shading_elm)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def set_page_border(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Size of the border
            border.set(qn('w:space'), '24')  # Space between the text and the border
            border.set(qn('w:color'), '000000')  # Black color
            pgBorders.append(border)

        sectPr.append(pgBorders)


def set_row_height(row, height):
    """
    Set the height of a table row.

    :param row: The table row to set the height.
    :param height: The desired height in points.
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def set_row_height_executive(row, height):
    """
    Set the height of a table row.

    :param row: The table row whose height will be set.
    :param height: The height value.
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trPr.append(trHeight)

def set_narrow_margins(doc, margin=0.5):
    for section in doc.sections:
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)


def set_document_table_cell_text_content(cell, text):
    """
    Sets the text and style for a cell in the content table.
    """
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(14)

def set_table_header_style(cell, text):
    """
    Sets the style for the header cells in the table.
    """
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
    # Set background color for the header cell
    cell_fill = OxmlElement('w:shd')
    cell_fill.set(qn('w:val'), 'clear')
    cell_fill.set(qn('w:color'), 'auto')
    cell_fill.set(qn('w:fill'), 'a0c4e4')
    cell._element.get_or_add_tcPr().append(cell_fill)
def set_table_header_style_dch(cell, text):
    """
    Sets the style for the header cells in the table.
    """
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
    # Set background color for the header cell
    cell_fill = OxmlElement('w:shd')
    cell_fill.set(qn('w:val'), 'clear')
    cell_fill.set(qn('w:color'), 'auto')
    cell_fill.set(qn('w:fill'), 'FFFFFF')
    cell._element.get_or_add_tcPr().append(cell_fill)


def download_report_word(selected_vulnerabilities):
    doc = Document()
    doc.add_heading('Web Application Vulnerability Assessment Report', 0)

    # Set narrow margins
    set_narrow_margins(doc, margin=0.5)
    page_width = 8.5 - 2 * 0.5  # Page width minus left and right margins


    # Set page border
    set_page_border(doc)
    # Add header with logo and CONFIDENTIAL text
    section = doc.sections[0]
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(8.5))
    header_table.autofit = True
    header_table.columns[0].width = Inches(4.25)
    header_table.columns[1].width = Inches(3.25)

    # Add logo to the first cell
    logo_cell = header_table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]
    run = logo_paragraph.add_run()
    logo_path = './static/assets/img/download.png'  # Update this with the correct path to your logo
    run.add_picture(logo_path, width=Inches(1))  # Adjust the width as necessary

    # Add CONFIDENTIAL text to the second cell
    header_text_cell = header_table.cell(0, 1)
    header_text_paragraph = header_text_cell.paragraphs[0]
    header_text_paragraph.text = "CONFIDENTIAL"
    header_text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = header_text_paragraph.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True

    # Add footer with page number and other text
    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(8.5))
    footer_table.autofit = True
    footer_table.columns[0].width = Inches(4.25)
    footer_table.columns[1].width = Inches(3.15)

    # Add page number to the first cell
    page_number_cell = footer_table.cell(0, 0)
    page_number_paragraph = page_number_cell.paragraphs[0]
    page_number_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = page_number_paragraph.add_run()
    run.text = "Page "
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')  # sets attribute on element
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)

    # Add footer text to the second cell
    footer_text_cell = footer_table.cell(0, 1)
    footer_text_paragraph = footer_text_cell.paragraphs[0]
    footer_text_paragraph.text = "CERT-In Audit report format: Version 1.0"
    footer_text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Set paragraph spacing to zero
    footer_text_paragraph.space_before = Pt(0)
    footer_text_paragraph.space_after = Pt(0)

    # Set cell padding and margins to zero
    tc = footer_text_cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    tcMar.set(qn('w:top'), "0")
    tcMar.set(qn('w:bottom'), "0")
    tcMar.set(qn('w:left'), "0")
    tcMar.set(qn('w:right'), "0")
    tcPr.append(tcMar)

    run = footer_text_paragraph.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True

    # Add the Document Preparation data as a table with heading
    if document_preparation_data:
        doc.add_paragraph().add_run().add_break()
        heading = doc.add_heading(level=1)
        # run = heading.add_run('Document Preparation')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(12)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=len(document_preparation_data) + 2, cols=2)  # Add an extra row for the heading
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Set the column widths to 1.8 and 5.8
        for row in table.rows:
            row.cells[0].width = Inches(1.8)
            row.cells[1].width = Inches(5.8)

        # Heading row
        heading_cells = table.rows[0].cells
        heading_cells[0].merge(heading_cells[1])  # Merge the two cells for the heading
        set_table_header_style(heading_cells[0], 'Document Preparation')
        heading_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        # Data rows
        for i, (key, value) in enumerate(document_preparation_data.items(), start=2):
            row_cells = table.rows[i].cells
            row_cells[0].text = key
            set_document_table_cell_text(row_cells[0], key)
            row_cells[1].text = str(value)
            set_document_table_cell_text(row_cells[1], str(value))
            row_color = 'e0ecf4' if i % 2 == 0 else 'FFFFFF'
            for cell in row_cells:
                color_cell(cell, row_color)
            # Add empty space between tables
        doc.add_paragraph().add_run().add_break()

    # Add the Document Change History data as a table with heading
    if document_change_history_data:
        doc.add_paragraph().add_run().add_break()

        # Create a table with an extra row for the heading
        table = doc.add_table(rows=len(document_change_history_data) + 2, cols=3)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Heading row
        heading_cells = table.rows[0].cells
        heading_cells[0].merge(heading_cells[1]).merge(heading_cells[2])  # Merge the three cells for the heading
        set_table_header_style(heading_cells[0], 'Document Change History')
        heading_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[1].cells
        set_table_header_style_dch(header_cells[0], 'Version')
        set_table_header_style_dch(header_cells[1], 'Date')
        set_table_header_style_dch(header_cells[2], 'Remarks')

        # Data rows
        for i, row_data in enumerate(document_change_history_data, start=2):
            row_cells = table.rows[i].cells
            row_cells[0].text = row_data.get('version', '')
            set_document_table_cell_text(row_cells[0], row_data.get('version', ''))
            row_cells[1].text = row_data.get('date', '')
            set_document_table_cell_text(row_cells[1], row_data.get('date', ''))
            row_cells[2].text = row_data.get('remarks', '')
            set_document_table_cell_text(row_cells[2], row_data.get('remarks', ''))
            row_color = 'e0ecf4' if i % 2 == 0 else 'FFFFFF'
            for cell in row_cells:
                color_cell(cell, row_color)

        # Add empty space between tables
        doc.add_paragraph().add_run().add_break()

    # Add the Document Distribution List data as a table with heading
    if document_distribution_list_data:
        doc.add_paragraph().add_run().add_break()
        # heading = doc.add_heading(level=1)
        # # heading_run = heading.add_run('Document Distribution List')
        # font = heading_run.font
        # font.name = 'Calibri'
        # font.size = Pt(20)
        # font.bold = True
        # heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=len(document_distribution_list_data) + 2, cols=4)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Heading row
        heading_row = table.rows[0].cells
        heading_row[0].merge(heading_row[1]).merge(heading_row[2]).merge(
            heading_row[3])  # Merge all four cells for the heading
        set_table_header_style(heading_row[0], 'Document Distribution List')
        heading_row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[1].cells
        set_table_header_style_dch(header_cells[0], 'Name')
        set_table_header_style_dch(header_cells[1], 'Organization')
        set_table_header_style_dch(header_cells[2], 'Designation')
        set_table_header_style_dch(header_cells[3], 'Email Id')

        # Data rows
        for i, row_data in enumerate(document_distribution_list_data, start=2):
            row_cells = table.rows[i].cells
            set_document_table_cell_text(row_cells[0], row_data.get('name', ''))
            set_document_table_cell_text(row_cells[1], row_data.get('organization', ''))
            set_document_table_cell_text(row_cells[2], row_data.get('designation', ''))
            set_document_table_cell_text(row_cells[3], row_data.get('email', ''))
            row_color = 'e0ecf4' if i % 2 == 0 else 'FFFFFF'
            for cell in row_cells:
                color_cell(cell, row_color)

        # Add empty space between tables
        doc.add_paragraph().add_run().add_break()

    # Add the Contents HTML data before vulnerabilities and on a new page
    if contents_data:
        doc.add_page_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('Contents')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=len(contents_data) + 1, cols=2)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Set white borders for the content table
        set_table_border_white(table)

        # Set the column widths to 6.8 and 0.8
        for row in table.rows:
            row.cells[0].width = Inches(7.2)
            row.cells[1].width = Inches(0.4)


        # Data rows
        for i, row_data in enumerate(contents_data, start=1):
            row_cells = table.rows[i].cells
            set_document_table_cell_text_content(row_cells[0], row_data.get('title', ''))
            set_document_table_cell_text_content(row_cells[1], row_data.get('page', ''))


    # Add the Introduction data before vulnerabilities and after Contents table, on a new page
    if document_data:
        doc.add_page_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('Introduction')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(document_data.get('introduction', ''))

    # Engagement Scope
    if engagement_scope_data:
        doc.add_page_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('Engagement Scope')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=len(engagement_scope_data) + 1, cols=10)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Header row
        header_cells = table.rows[0].cells
        headers = [
            'S. No', 'Asset Description', 'Criticality of Asset', 'Internal IP Address',
            'URL', 'Public IP Address', 'Location', 'Hash Value (in case of applications)',
            'Version (in case of applications)',
            'Other details such as make and model in case of network devices or security devices'
        ]
        for i, header in enumerate(headers):
            set_header_style_engagement(header_cells[i], header, font_size=8)

        # Set the header row height
        set_row_height_executive(table.rows[0], 2000)

        # Set the column widths
        column_widths = [
            0.1, 0.7, 0.7, 0.6, 0.7, 0.7, 0.7, 0.7, 0.6, 1.3
        ]
        for row in table.rows:
            for i, width in enumerate(column_widths):
                row.cells[i].width = Inches(width)

        # Data rows
        for i, row_data in enumerate(engagement_scope_data, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = row_data.get('s_no', '')
            row_cells[1].text = row_data.get('asset_description', '')
            row_cells[2].text = row_data.get('criticality_of_asset', '')
            row_cells[3].text = row_data.get('internal_ip_address', '')
            row_cells[4].text = row_data.get('url', '')
            row_cells[5].text = row_data.get('public_ip_address', '')
            row_cells[6].text = row_data.get('location', '')
            row_cells[7].text = row_data.get('hash_value', '')
            row_cells[8].text = row_data.get('version', '')
            row_cells[9].text = row_data.get('other_details', '')
            row_color = 'e0ecf4' if i % 2 == 0 else 'FFFFFF'
            for cell in row_cells:
                color_cell(cell, row_color)

            for cell in row_cells:
                set_document_cell_text_engagement(cell, cell.text)
    # Add the Details of the Auditing Team table after Engagement Scope and before vulnerabilities
    if auditing_team_data:
        doc.add_paragraph().add_run().add_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('Details of the Auditing Team')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=len(auditing_team_data) + 1, cols=6)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Header row
        header_cells = table.rows[0].cells
        headers = [
            'S. No', 'Name', 'Designation', 'Email Id',
            'Professional Qualifications/ Certifications',
            'Whether the resource has Been listed in the Snapshot information published on CERT-In’s website (Yes/No)'
        ]
        for i, header in enumerate(headers):
            set_header_style_Audit(header_cells[i], header, font_size=8)

        # Data rows
        for i, row_data in enumerate(auditing_team_data, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = row_data.get('s_no', '')
            row_cells[1].text = row_data.get('name', '')
            row_cells[2].text = row_data.get('designation', '')
            row_cells[3].text = row_data.get('email', '')
            row_cells[4].text = row_data.get('qualifications', '')
            row_cells[5].text = row_data.get('listed', '')
            row_color = 'e0ecf4' if i % 2 == 0 else 'FFFFFF'
            for cell in row_cells:
                color_cell(cell, row_color)

            for cell in row_cells:
                set_document_table_cell_text(cell, cell.text)
     # Add the Audit Activities and Timelines table after Engagement Scope and before vulnerabilities
    if audit_timeline_data:
        heading = doc.add_heading(level=1)
        run = heading.add_run('Audit Activities and Timelines')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=len(audit_timeline_data) + 1, cols=4)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Header row
        header_cells = table.rows[0].cells
        set_table_header_style_AA(header_cells[0], 'Sl. No.')
        set_table_header_style_AA(header_cells[1], 'Activities')
        set_table_header_style_AA(header_cells[2], 'Start Date')
        set_table_header_style_AA(header_cells[3], 'End Date')

        # Data rows
        for i, row_data in enumerate(audit_timeline_data, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = row_data.get('s_no', '')
            set_document_table_cell_text(row_cells[0], row_data.get('s_no', ''))
            row_cells[1].text = row_data.get('activities', '')
            set_document_table_cell_text(row_cells[1], row_data.get('activities', ''))
            row_cells[2].text = row_data.get('start_date', '')
            set_document_table_cell_text(row_cells[2], row_data.get('start_date', ''))
            row_cells[3].text = row_data.get('end_date', '')
            set_document_table_cell_text(row_cells[3], row_data.get('end_date', ''))
            row_color = 'e0ecf4' if i % 2 == 0 else 'FFFFFF'
            for cell in row_cells:
                color_cell(cell, row_color)

        # Add empty space between tables
        doc.add_paragraph().add_run().add_break()

    # Add the Audit Methodology and Criteria section before vulnerabilities, on a new page
    if audit_methodology_criteria_data:
        doc.add_page_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('Audit Methodology and Criteria/ Standard referred for audit')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the paragraph content
        methodology_text = audit_methodology_criteria_data.get('methodology', '')
        doc.add_paragraph(methodology_text)

    # Add the OWASP Top 2021 Vulnerabilities table
    if owasp_data:
        doc.add_paragraph().add_run().add_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('OWASP Top 2021 Ten Most Critical Web Application Vulnerabilities Mapping')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=len(owasp_data) + 1, cols=3)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Set column widths for OWASP table
        for row in table.rows:
            row.cells[0].width = Inches(0.2)
            row.cells[1].width = Inches(5.3)
            row.cells[2].width = Inches(2.0)

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Sl. No.', 'Security Risk', 'Present in Web Application']
        for i, header in enumerate(headers):
            set_table_header_style_owasp(header_cells[i], header)

        # Data rows
        for i, row_data in enumerate(owasp_data, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(row_data.get('s_no', ''))
            row_cells[1].text = row_data.get('security_risk', '')
            present_value = row_data.get('present', '')
            row_color = 'e0ecf4' if i % 2 == 0 else 'FFFFFF'
            for cell in row_cells:
                color_cell(cell, row_color)

            # Map "No" and "Yes" to their full descriptions if necessary
            if present_value == "No":
                present_value = "No (Examined & Not Found)"
            elif present_value == "Yes":
                present_value = "Yes (Found)"

            row_cells[2].text = present_value

            for cell in row_cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add the SANS-25 Vulnerabilities table
    if sans_data:
        doc.add_paragraph().add_run().add_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('SANS-25 Most Critical Web Application Vulnerabilities Mapping')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=len(sans_data) + 1, cols=3)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Set column widths for SANS table
        for row in table.rows:
            row.cells[0].width = Inches(0.2)
            row.cells[1].width = Inches(5.3)
            row.cells[2].width = Inches(2.0)

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Sl. No.', 'Security Risk', 'Present in Web Application']
        for i, header in enumerate(headers):
            set_table_header_style_sans(header_cells[i], header)

        # Data rows
        for i, row_data in enumerate(sans_data, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(row_data.get('s_no', ''))
            row_cells[1].text = row_data.get('security_risk', '')
            present_value = row_data.get('present', '')
            row_color = 'e0ecf4' if i % 2 == 0 else 'FFFFFF'
            for cell in row_cells:
                color_cell(cell, row_color)

            # Map "No" and "Yes" to their full descriptions if necessary
            if present_value == "No":
                present_value = "No (Examined & Not Found)"
            elif present_value == "Yes":
                present_value = "Yes (Found)"

            row_cells[2].text = present_value

            for cell in row_cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # Add the Tools/Software Used table before vulnerabilities
    if tools_data:
        doc.add_paragraph().add_run().add_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('Tools/Software Used')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=len(tools_data) + 1, cols=4)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Header row
        header_cells = table.rows[0].cells
        set_table_header_style_dch(header_cells[0], 'Sl. No.')
        set_table_header_style_dch(header_cells[1], 'Name of Tool/Software Used')
        set_table_header_style_dch(header_cells[2], 'Version of the Tool/Software Used')
        set_table_header_style_dch(header_cells[3], 'Open Source/Licensed')

        # Data rows
        for i, row_data in enumerate(tools_data, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(row_data.get('serialNumber', ''))
            set_document_table_cell_text(row_cells[0], str(row_data.get('serialNumber', '')))
            row_cells[1].text = row_data.get('toolName', '')
            set_document_table_cell_text(row_cells[1], row_data.get('toolName', ''))
            row_cells[2].text = row_data.get('toolVersion', '')
            set_document_table_cell_text(row_cells[2], row_data.get('toolVersion', ''))
            row_cells[3].text = row_data.get('toolType', '')
            set_document_table_cell_text(row_cells[3], row_data.get('toolType', ''))
            row_color = 'e0ecf4' if i % 2 == 0 else 'FFFFFF'
            for cell in row_cells:
                color_cell(cell, row_color)

        # Add empty space between tables
        doc.add_paragraph().add_run().add_break()

    # Executive Summary
    if selected_vulnerabilities:
        doc.add_page_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('Executive Summary')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Define the table structure
        table = doc.add_table(rows=len(selected_vulnerabilities) + 1, cols=11)
        table.style = 'Table Grid'
        set_table_full_width(table, page_width)  # Set the table to full page width

        # Header row
        headers = [
            'S. No ', 'Affected Asset i.e. IP/URL/Application etc', 'Observation/ Vulnerability title',
            'CVE/CWE ', 'Control Objective #', 'Control Name #', 'Audit Requirement #',
            'Severity', 'Recommendation', 'Reference', 'New or Repeat observation'
        ]

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ""
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(header)
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(9)
            font.bold = True
            font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # Set font color to #FFFFFF
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '31473A')  # Set background color to #31473A
            cell._element.get_or_add_tcPr().append(shading_elm)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Set vertical text direction
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            text_direction = OxmlElement('w:textDirection')
            text_direction.set(qn('w:val'), 'btLr')
            tcPr.append(text_direction)

        # Set the header row height
        set_row_height_executive(table.rows[0], 2000)

        # Data rows
        for i, vuln in enumerate(selected_vulnerabilities, start=1):
            cells = table.rows[i].cells
            cells[0].text = str(i)
            cells[1].text = vuln.get('assets', 'N/A')
            cells[2].text = vuln.get('name', 'N/A')
            cells[3].text = vuln.get('cve_cwe', 'N/A')
            cells[4].text = vuln.get('control_objective', 'N/A')
            cells[5].text = vuln.get('control_name', 'N/A')
            cells[6].text = vuln.get('audit_requirement', 'N/A')
            cells[7].text = vuln.get('severity', 'N/A')
            cells[8].text = vuln.get('recommendation', 'N/A')
            cells[9].text = vuln.get('severity', 'N/A')
            cells[10].text = vuln.get('new_repeat_observation', 'N/A')

            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(11)
                        font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Set font color to black
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add the Vulnerabilities data as tables with headings, each starting on a new page
        for vuln in selected_vulnerabilities:
            doc.add_page_break()
            heading = doc.add_heading(level=1)
            run = heading.add_run(f'Vulnerability: {vuln.get("name", "N/A")}')
            font = run.font
            font.name = 'Calibri'
            font.size = Pt(14)
            font.bold = True
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            headings = [
                ('Affected Asset:', vuln.get('assets', 'N/A')),
                ('Vulnerability Title:', vuln.get('name', 'N/A')),
                ('Business Impact:', vuln.get('business_impact', 'N/A')),
                ('Detailed observation:', vuln.get('detailed_observation', 'N/A')),
                ('CVE/CWE:', vuln.get('cve_cwe', 'N/A')),
                ('Proof of Concept:', ''),  # Placeholder for POC images and descriptions
                ('Vulnerable Points:', vuln.get('vulnerable_points', 'N/A')),
                ('Recommendation:', vuln.get('recommendation', 'N/A')),
                ('Reference:', vuln.get('reference', 'N/A')),
                ('New or Repeat observation:', vuln.get('new_repeat_observation', 'N/A'))
            ]

            table = doc.add_table(rows=len(headings), cols=2)
            table.style = 'Table Grid'
            set_table_full_width(table, page_width)  # Set the table to full page width

            for row in table.rows:
                row.cells[0].width = Inches(7.5)
                row.cells[1].width = Inches(0.3)

            severity_color_map = {
                'informational': '00FF00',  # Green
                'low': 'FFFF00',  # Yellow
                'medium': 'FFA500',  # Orange
                'high': 'FF0000'  # Red
            }

            severity = vuln.get('severity', 'N/A').lower()
            color = severity_color_map.get(severity, 'FFFFFF')  # Default to white if severity is not recognized

            for i, (heading, content) in enumerate(headings):
                if heading == 'Proof of Concept:':
                    set_vulnerability_table_cell_text(table.cell(i, 0), heading, '', 10, bold_heading=True)
                    color_cell(table.cell(i, 1), color)
                    for poc_file, poc_description in vuln.get('pocs', []):
                        paragraph = table.cell(i, 0).add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(os.path.join(app.config['UPLOAD_FOLDER'], poc_file), width=Inches(5))
                        table.cell(i, 0).add_paragraph(poc_description)
                else:
                    set_vulnerability_table_cell_text(table.cell(i, 0), heading, content, 10, bold_heading=True)
                    color_cell(table.cell(i, 1), color)
    # Add the Acronym section before vulnerabilities and after any previous sections, on a new page
    if 'acronym' in disclaimer_data:
        doc.add_page_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('Acronym')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(disclaimer_data['acronym'])

    # Add the Disclaimer section before vulnerabilities and after any previous sections, on a new page
    if 'disclaimer' in disclaimer_data:
        doc.add_paragraph().add_run().add_break()
        heading = doc.add_heading(level=1)
        run = heading.add_run('Disclaimer')
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        font.bold = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(disclaimer_data['disclaimer'])

    # Save the document for demonstration
    report_file = io.BytesIO()
    doc.save(report_file)
    report_file.seek(0)

    return send_file(report_file,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name='vulnerabilities_report.docx')


def set_table_full_width(table, page_width=7.5):
    num_columns = len(table.columns)
    column_width = page_width / num_columns
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(column_width)
# Function to set white table borders
def set_table_border_white(table):
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')

    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_element = OxmlElement(f'w:{border}')
        border_element.set(qn('w:val'), 'single')
        border_element.set(qn('w:sz'), '8')
        border_element.set(qn('w:space'), '0')
        border_element.set(qn('w:color'), 'FFFFFF')  # Set border color to white
        tblBorders.append(border_element)

    tbl.tblPr.append(tblBorders)


def set_table_header_style(cell, text):
    cell.text = text
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.bold = True
    cell_font.size = Pt(12)
    cell_font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Header text color
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell._element.get_or_add_tcPr().append(create_shading_element('a0c4e4'))


def set_document_table_cell_text(cell, text):
    cell.text = text
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(10)
    cell_font.bold = False



def create_shading_element(color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    return shading_elm

def color_cell(cell, color):
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = create_shading_element(color)
    cell_properties.append(cell_shading)


@app.route('/discard_report', methods=['POST'])
def discard_report():
    global selected_vulnerabilities, document_data
    selected_vulnerabilities = []
    document_data = {}
    return jsonify({'message': 'Report discarded successfully.'})

@app.route('/redirect_to_vulnerabilities', methods=['POST'])
def redirect_to_vulnerabilities():
    return jsonify({'success': True, 'redirect_url': url_for('show_vulnerabilities', category='N/A')})


@app.route('/document_distribution_list', methods=['GET', 'POST'])
def document_distribution_list():
    if request.method == 'POST':
        global document_distribution_list_data
        document_distribution_list_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})

    return render_template('Document Distribution List.html', document_data=document_distribution_list_data)

@app.route('/save_document_distribution_list', methods=['POST'])
def save_document_distribution_list():
    global document_distribution_list_data
    document_distribution_list_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_document_distribution_list_data', methods=['GET'])
def get_document_distribution_list_data():
    return jsonify(document_distribution_list_data)


@app.route('/contents', methods=['GET', 'POST'])
def contents():
    if request.method == 'POST':
        global contents_data
        contents_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})

    return render_template('Content.html', document_data=contents_data)

@app.route('/save_contents', methods=['POST'])
def save_contents():
    global contents_data
    contents_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_contents_data', methods=['GET'])
def get_contents_data():
    return jsonify(contents_data)



@app.route('/introduction', methods=['GET', 'POST'])
def introduction():
    if request.method == 'POST':
        global document_data
        document_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})

    return render_template('Introduction.html', document_data=document_data)

@app.route('/save_introduction', methods=['POST'])
def save_introduction():
    global document_data
    document_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_introduction_data', methods=['GET'])
def get_introduction_data():
    return jsonify(document_data)

@app.route('/engagement_scope', methods=['GET', 'POST'])
def engagement_scope():
    if request.method == 'POST':
        global engagement_scope_data
        engagement_scope_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})

    return render_template('Engagement Scope.html', engagement_scope_data=engagement_scope_data)

@app.route('/save_engagement_scope', methods=['POST'])
def save_engagement_scope():
    global engagement_scope_data
    engagement_scope_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_engagement_scope_data', methods=['GET'])
def get_engagement_scope_data():
    return jsonify(engagement_scope_data)

@app.route('/details_of_the_auditing_team', methods=['GET', 'POST'])
def details_of_the_auditing_team():
    if request.method == 'POST':
        global auditing_team_data
        auditing_team_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})

    return render_template('Details of Auditing Team.html', auditing_team_data=auditing_team_data)

@app.route('/save_auditing_team', methods=['POST'])
def save_auditing_team():
    global auditing_team_data
    auditing_team_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_auditing_team_data', methods=['GET'])
def get_auditing_team_data():
    return jsonify(auditing_team_data)



@app.route('/audit_activities_and_timeline', methods=['GET', 'POST'])
def audit_activities_and_timeline():
    global audit_timeline_data
    if request.method == 'POST':
        audit_timeline_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})

    return render_template('Audit Activities and Timeline.html', audit_timeline_data=audit_timeline_data)

@app.route('/save_audit_timeline', methods=['POST'])
def save_audit_timeline():
    global audit_timeline_data
    audit_timeline_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_audit_timeline_data', methods=['GET'])
def get_audit_timeline_data():
    return jsonify(audit_timeline_data)

@app.route('/audit_methodology_and_criteria', methods=['GET', 'POST'])
def audit_methodology_and_criteria():
    if request.method == 'POST':
        global audit_methodology_criteria_data
        audit_methodology_criteria_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})

    return render_template('Audit Methodology and Criteria.html', audit_methodology_criteria_data=audit_methodology_criteria_data)

@app.route('/save_audit_methodology_and_criteria', methods=['POST'])
def save_audit_methodology_and_criteria():
    global audit_methodology_criteria_data
    audit_methodology_criteria_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_audit_methodology_and_criteria_data', methods=['GET'])
def get_audit_methodology_and_criteria_data():
    return jsonify(audit_methodology_criteria_data)


@app.route('/owasp', methods=['GET', 'POST'])
def owasp():
    global owasp_data
    if request.method == 'POST':
        owasp_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})
    return render_template('OWASP Top 2021 Ten Most Critical Web Application Vulnerabilities Mapping.html', owasp_data=owasp_data or [])

@app.route('/save_owasp_vulnerabilities', methods=['POST'])
def save_owasp_vulnerabilities():
    global owasp_data
    owasp_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_owasp_data', methods=['GET'])
def get_owasp_data():
    return jsonify(owasp_data)


@app.route('/sans', methods=['GET', 'POST'])
def sans():
    global sans_data
    if request.method == 'POST':
        sans_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})
    # Initialize sans_data if not already defined
    if 'sans_data' not in globals():
        sans_data = []
    return render_template('SANS-25 Most Critical Web Application Vulnerabilities Mapping.html', sans_data=sans_data)

@app.route('/save_sans_vulnerabilities', methods=['POST'])
def save_sans_vulnerabilities():
    global sans_data
    sans_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_sans_data', methods=['GET'])
def get_sans_data():
    return jsonify(sans_data)


@app.route('/tools_and_software_used', methods=['GET', 'POST'])
def tools_and_software_used():
    global tools_data
    if request.method == 'POST':
        tools_data = request.json
        return jsonify({'success': True, 'message': 'Data saved successfully'})
    return render_template('Tools and Software Used.html', tools_data=tools_data)

@app.route('/save_tools_data', methods=['POST'])
def save_tools_data():
    global tools_data
    tools_data = request.json
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_tools_data', methods=['GET'])
def get_tools_data():
    return jsonify(tools_data)


@app.route('/clear_data', methods=['POST'])
def clear_data():
    # Your logic to clear data
    return 'Data cleared', 200


# Routes for Disclaimer and Acronym
@app.route('/disclaimer', methods=['GET', 'POST'])
def disclaimer():
    if request.method == 'POST':
        global disclaimer_data
        disclaimer_data['acronym'] = request.json.get('acronym', '')
        disclaimer_data['disclaimer'] = request.json.get('disclaimer', '')
        return jsonify({'success': True, 'message': 'Data saved successfully'})
    return render_template('Disclaimer.html', disclaimer_data=disclaimer_data)


@app.route('/save_disclaimer', methods=['POST'])
def save_disclaimer():
    global disclaimer_data
    disclaimer_data['acronym'] = request.json.get('acronym', '')
    disclaimer_data['disclaimer'] = request.json.get('disclaimer', '')
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})


@app.route('/get_disclaimer_data', methods=['GET'])
def get_disclaimer_data():
    return jsonify(disclaimer_data)


@app.route('/appendix', methods=['GET', 'POST'])
def appendix():
    if request.method == 'POST':
        global appendix_data
        appendix_data['risk_ranking_approach'] = request.json.get('risk_ranking_approach', '')
        appendix_data['likelihood'] = request.json.get('likelihood', '')
        appendix_data['impact'] = request.json.get('impact', '')
        return jsonify({'success': True, 'message': 'Data saved successfully'})
    return render_template('Appendix.html', appendix_data=appendix_data)

@app.route('/save_appendix', methods=['POST'])
def save_appendix():
    global appendix_data
    appendix_data['risk_ranking_approach'] = request.json.get('risk_ranking_approach', '')
    appendix_data['likelihood'] = request.json.get('likelihood', '')
    appendix_data['impact'] = request.json.get('impact', '')
    return jsonify({'success': True, 'message': 'Data saved successfully', 'redirect_url': url_for('document_control')})

@app.route('/get_appendix_data', methods=['GET'])
def get_appendix_data():
    return jsonify(appendix_data)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True, threaded=True)
